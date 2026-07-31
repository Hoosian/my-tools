"""Core logic for removing strikethrough runs from Word documents."""

from docx import Document
from docx.table import Table


def _is_strikethrough(run) -> bool:
    """Return True if the run has direct strikethrough formatting."""
    return bool(run.font.strike) or bool(run.font.double_strike)


def _remove_strikethrough_from_paragraph(paragraph) -> None:
    """Remove all strikethrough runs from a single paragraph."""
    for run in list(paragraph.runs):
        if _is_strikethrough(run):
            paragraph._p.remove(run._r)


def _remove_strikethrough_from_paragraphs(paragraphs) -> None:
    for paragraph in paragraphs:
        _remove_strikethrough_from_paragraph(paragraph)


def _remove_strikethrough_from_table(table: Table) -> None:
    """Recursively remove strikethrough runs from a table and nested tables."""
    for row in table.rows:
        for cell in row.cells:
            _remove_strikethrough_from_paragraphs(cell.paragraphs)
            for nested_table in cell.tables:
                _remove_strikethrough_from_table(nested_table)


def _remove_strikethrough_from_header_footer(section) -> None:
    """Process all headers and footers in a section."""
    headers = [section.header]
    footers = [section.footer]

    if section.different_first_page_header_footer:
        headers.extend([section.first_page_header])
        footers.extend([section.first_page_footer])

    if getattr(section, "different_odd_even_pages_header_footer", False):
        headers.extend([section.even_page_header])
        footers.extend([section.even_page_footer])

    for header in headers:
        if header is None:
            continue
        _remove_strikethrough_from_paragraphs(header.paragraphs)
        for table in header.tables:
            _remove_strikethrough_from_table(table)

    for footer in footers:
        if footer is None:
            continue
        _remove_strikethrough_from_paragraphs(footer.paragraphs)
        for table in footer.tables:
            _remove_strikethrough_from_table(table)


def remove_strikethrough(input_path: str, output_path: str) -> None:
    """Remove all direct strikethrough runs and save a new document."""
    doc = Document(input_path)

    # Body paragraphs and tables
    _remove_strikethrough_from_paragraphs(doc.paragraphs)
    for table in doc.tables:
        _remove_strikethrough_from_table(table)

    # Headers and footers
    for section in doc.sections:
        _remove_strikethrough_from_header_footer(section)

    doc.save(output_path)
