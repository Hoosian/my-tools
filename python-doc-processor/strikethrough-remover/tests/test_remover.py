import os
import tempfile
import unittest

from docx import Document

from strike_remover.remover import remove_strikethrough


class TestStrikethroughRemover(unittest.TestCase):
    @staticmethod
    def _create_fixture(path: str) -> None:
        doc = Document()

        # Paragraph with mixed normal and strikethrough runs
        p = doc.add_paragraph()
        p.add_run("Keep ")
        p.add_run("removed").font.strike = True
        p.add_run(" text")

        # Paragraph with only normal text
        doc.add_paragraph("All normal text here")

        # Paragraph with double strikethrough
        p2 = doc.add_paragraph()
        p2.add_run("Double ")
        p2.add_run("gone").font.double_strike = True

        # Table with a strikethrough cell (cell becomes empty but row is kept)
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "normal cell"
        cell = table.cell(0, 1)
        cell.text = ""
        cell.paragraphs[0].add_run("strike cell").font.strike = True

        # Table with a fully empty row after cleanup (should be removed)
        table2 = doc.add_table(rows=2, cols=2)
        table2.cell(0, 0).text = "keep row 1"
        table2.cell(0, 1).text = "keep row 2"
        table2.cell(1, 0).paragraphs[0].add_run("strike1").font.strike = True
        table2.cell(1, 1).paragraphs[0].add_run("strike2").font.strike = True

        doc.save(path)

    def test_removes_strikethrough_and_double_strike(self):
        with tempfile.TemporaryDirectory() as tmp:
            input_path = os.path.join(tmp, "input.docx")
            output_path = os.path.join(tmp, "output.docx")
            self._create_fixture(input_path)

            remove_strikethrough(input_path, output_path)
            doc = Document(output_path)

            paragraphs = [p.text for p in doc.paragraphs]

            self.assertEqual(paragraphs[0], "Keep text")
            self.assertNotIn("removed", paragraphs[0])
            self.assertEqual(paragraphs[1], "All normal text here")
            self.assertEqual(paragraphs[2], "Double ")
            self.assertNotIn("gone", paragraphs[2])

            table = doc.tables[0]
            self.assertEqual(table.cell(0, 0).text, "normal cell")
            self.assertEqual(table.cell(0, 1).text, "")

            table2 = doc.tables[1]
            self.assertEqual(len(table2.rows), 1)
            self.assertEqual(table2.cell(0, 0).text, "keep row 1")
            self.assertEqual(table2.cell(0, 1).text, "keep row 2")


if __name__ == "__main__":
    unittest.main()
