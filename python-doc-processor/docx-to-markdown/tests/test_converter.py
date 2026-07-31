import os
import tempfile
import unittest

from docx import Document

from docx_to_markdown.converter import docx_to_markdown


class TestDocxToMarkdown(unittest.TestCase):
    @staticmethod
    def _create_fixture(path: str) -> None:
        doc = Document()

        doc.add_heading("Title", level=1)
        doc.add_heading("Subtitle", level=2)

        doc.add_paragraph("This is a normal paragraph.")

        p = doc.add_paragraph()
        p.add_run("Bold text").bold = True
        p.add_run(" and ")
        p.add_run("italic text").italic = True
        p.add_run(".")

        doc.add_paragraph("Bullet item 1", style="List Bullet")
        doc.add_paragraph("Bullet item 2", style="List Bullet")

        doc.add_paragraph("Numbered item 1", style="List Number")
        doc.add_paragraph("Numbered item 2", style="List Number")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"

        doc.add_paragraph("")
        doc.add_paragraph("Final paragraph.")

        doc.save(path)

    def test_converts_full_document(self):
        with tempfile.TemporaryDirectory() as tmp:
            input_path = os.path.join(tmp, "input.docx")
            output_path = os.path.join(tmp, "output.md")
            self._create_fixture(input_path)

            result = docx_to_markdown(input_path, output_path)
            with open(output_path, encoding="utf-8") as f:
                written = f.read()

            self.assertEqual(result, written)
            self.assertIn("# Title", result)
            self.assertIn("## Subtitle", result)
            self.assertIn("This is a normal paragraph.", result)
            self.assertIn("**Bold text**", result)
            self.assertIn("*italic text*", result)
            self.assertIn("- Bullet item 1", result)
            self.assertIn("1. Numbered item 1", result)
            self.assertIn("| Name | Age |", result)
            self.assertIn("| --- | --- |", result)
            self.assertIn("| Alice | 30 |", result)
            self.assertIn("Final paragraph.", result)


if __name__ == "__main__":
    unittest.main()
